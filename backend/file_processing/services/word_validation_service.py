import zipfile
import math
import xml.etree.ElementTree as ET
from io import BytesIO
from dataclasses import dataclass
from typing import Any, Optional, Tuple

from docx import Document

EXT_DOCX = ".docx"
EXT_DOC = ".doc"
MAX_WORD_PAGES = 100
DOCX_CHARS_PER_PAGE_ESTIMATE = 2500
DOCX_WORDS_PER_PAGE_ESTIMATE = 350
DOCX_BLOCKS_PER_PAGE_ESTIMATE = 35
DOCX_DEFAULT_PAGE_WIDTH_INCH = 8.5
DOCX_DEFAULT_PAGE_HEIGHT_INCH = 11.0
DOCX_DEFAULT_MARGIN_INCH = 1.0
DOCX_IMAGE_DENSE_THRESHOLD = 0.35
DOCX_IMAGE_PAGE_AREA_FACTOR = 0.65
DOCX_IMAGE_MEDIUM_AREA_INCH2 = 6.0
DOCX_IMAGE_SMALL_AREA_INCH2 = 1.5
WORD_CORRUPT_ERROR = "Word file is corrupt or has an invalid structure."
WORD_PROTECTED_ERROR = "Word file is password-protected."
OLE_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
WORD_DOCUMENT_XML = "word/document.xml"


@dataclass
class WordValidationContext:
    uploaded_file: Any
    ext: str
    page_count: int = 0


class WordValidationHandler:
    def __init__(self, next_handler: Optional["WordValidationHandler"] = None):
        self._next_handler = next_handler

    def set_next(
        self, next_handler: "WordValidationHandler"
    ) -> "WordValidationHandler":
        self._next_handler = next_handler
        return next_handler

    def _next(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        if self._next_handler is None:
            return True, None
        return self._next_handler.handle(context)

    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        raise NotImplementedError


class DocxEncryptedValidationHandler(WordValidationHandler):
    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        # Encrypted OOXML files are wrapped in OLE container, not regular ZIP-based DOCX.
        if is_ole_container(context.uploaded_file):
            return False, WORD_PROTECTED_ERROR
        return self._next(context)


class DocxStructureValidationHandler(WordValidationHandler):
    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        try:
            context.uploaded_file.seek(0)
            docx_bytes = context.uploaded_file.read()
            with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or WORD_DOCUMENT_XML not in names:
                    return False, WORD_CORRUPT_ERROR

                context.page_count = extract_docx_page_count(archive, docx_bytes)
        except Exception:
            return False, WORD_CORRUPT_ERROR
        finally:
            context.uploaded_file.seek(0)

        return self._next(context)


class DocEncryptedValidationHandler(WordValidationHandler):
    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        if not is_ole_container(context.uploaded_file):
            return False, WORD_CORRUPT_ERROR

        try:
            context.uploaded_file.seek(0)
            head = context.uploaded_file.read(4096)
            context.uploaded_file.seek(0)
            if b"EncryptedPackage" in head or b"EncryptionInfo" in head:
                return False, WORD_PROTECTED_ERROR
        except Exception:
            return False, WORD_CORRUPT_ERROR

        return self._next(context)


class DocStructureValidationHandler(WordValidationHandler):
    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        if not is_ole_container(context.uploaded_file):
            return False, WORD_CORRUPT_ERROR

        try:
            context.uploaded_file.seek(0)
            content = context.uploaded_file.read(1024 * 1024)
            context.uploaded_file.seek(0)

            if b"WordDocument" not in content:
                return False, WORD_CORRUPT_ERROR

            context.page_count = estimate_doc_page_count(content)
        except Exception:
            return False, WORD_CORRUPT_ERROR

        return self._next(context)


class WordPageCountValidationHandler(WordValidationHandler):
    def handle(self, context: WordValidationContext) -> Tuple[bool, Optional[str]]:
        if context.page_count > MAX_WORD_PAGES:
            return (
                False,
                f"Word exceeds the maximum allowed page count of {MAX_WORD_PAGES}.",
            )
        return True, None


def _build_word_validation_chain(ext: str) -> Optional[WordValidationHandler]:
    if ext == EXT_DOCX:
        start = DocxEncryptedValidationHandler()
        start.set_next(DocxStructureValidationHandler()).set_next(
            WordPageCountValidationHandler()
        )
        return start

    if ext == EXT_DOC:
        start = DocEncryptedValidationHandler()
        start.set_next(DocStructureValidationHandler()).set_next(
            WordPageCountValidationHandler()
        )
        return start

    return None


def validate_word(uploaded_file: Any, ext: str) -> Tuple[bool, Optional[str]]:
    chain = _build_word_validation_chain(ext)
    if chain is None:
        return False, "Unsupported file type."

    context = WordValidationContext(uploaded_file=uploaded_file, ext=ext)
    return chain.handle(context)


def check_docx_encrypted(uploaded_file: Any) -> Tuple[bool, Optional[str]]:
    # Encrypted OOXML files are wrapped in OLE container, not regular ZIP-based DOCX.
    if is_ole_container(uploaded_file):
        return False, WORD_PROTECTED_ERROR
    return True, None


def check_docx_structure(uploaded_file: Any) -> Tuple[bool, Any]:
    try:
        uploaded_file.seek(0)
        docx_bytes = uploaded_file.read()
        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            names = set(archive.namelist())
            if "[Content_Types].xml" not in names or WORD_DOCUMENT_XML not in names:
                return False, WORD_CORRUPT_ERROR

            page_count = extract_docx_page_count(archive, docx_bytes)
            return True, page_count
    except Exception:
        return False, WORD_CORRUPT_ERROR
    finally:
        uploaded_file.seek(0)


def extract_docx_page_count(
    archive: zipfile.ZipFile, docx_bytes: Optional[bytes] = None
) -> int:
    app_pages = 0
    try:
        app_xml = archive.read("docProps/app.xml")
        root = ET.fromstring(app_xml)
        for element in root.iter():
            if element.tag.endswith("Pages") and element.text:
                app_pages = max(int(element.text), 0)
                if app_pages > 0:
                    break
    except Exception:
        pass

    try:
        document_xml = archive.read(WORD_DOCUMENT_XML)
    except Exception:
        return app_pages

    estimated_pages = _estimate_docx_pages_from_document_xml(document_xml)
    python_docx_pages = _estimate_docx_pages_with_python_docx(docx_bytes)
    return max(app_pages, estimated_pages, python_docx_pages)


def _estimate_docx_pages_with_python_docx(docx_bytes: Optional[bytes]) -> int:
    if not docx_bytes:
        return 0

    document = _load_docx_document(docx_bytes)
    if document is None:
        return 0

    usable_page_area = _estimate_docx_usable_page_area(document)
    paragraph_count, paragraph_word_count, explicit_page_breaks = (
        _collect_paragraph_metrics(document)
    )
    image_count, total_image_area = _collect_body_image_metrics(document)
    table_row_count, table_word_count = _collect_table_metrics(document)

    word_count = paragraph_word_count + table_word_count
    section_estimate = _estimate_section_pages(document)
    word_estimate = _estimate_pages_from_words(word_count)
    structure_estimate = _estimate_pages_from_blocks(paragraph_count + table_row_count)
    break_estimate = _estimate_pages_from_explicit_breaks(explicit_page_breaks)
    image_estimate = _estimate_pages_from_images(
        image_count=image_count,
        total_image_area=total_image_area,
        usable_page_area=usable_page_area,
    )

    return max(
        break_estimate,
        section_estimate,
        word_estimate,
        structure_estimate,
        image_estimate,
    )


def _load_docx_document(docx_bytes: bytes) -> Optional[Document]:
    try:
        return Document(BytesIO(docx_bytes))
    except Exception:
        return None


def _collect_paragraph_metrics(document: Document) -> Tuple[int, int, int]:
    paragraph_count = 0
    word_count = 0
    explicit_page_breaks = 0

    for paragraph in document.paragraphs:
        paragraph_count += 1
        word_count += _count_words(paragraph.text)
        explicit_page_breaks += _count_explicit_breaks_in_paragraph(paragraph)

    return paragraph_count, word_count, explicit_page_breaks


def _count_words(text: str) -> int:
    cleaned_text = (text or "").strip()
    if not cleaned_text:
        return 0
    return len(cleaned_text.split())


def _count_explicit_breaks_in_paragraph(paragraph: Any) -> int:
    breaks = 0
    for node in paragraph._p.iter():
        node_name = _local_name(node.tag)
        if node_name == "lastRenderedPageBreak":
            breaks += 1
            continue
        if node_name == "br" and _is_page_break_node(node):
            breaks += 1
    return breaks


def _is_page_break_node(node: Any) -> bool:
    break_type = ""
    for key, value in node.attrib.items():
        if key.endswith("type"):
            break_type = (value or "").strip().lower()
            break
    return break_type == "page"


def _collect_body_image_metrics(document: Document) -> Tuple[int, float]:
    body_node = _safe_document_body(document)
    if body_node is None:
        return 0, 0.0

    image_count = 0
    total_image_area = 0.0
    for node in body_node.iter():
        node_name = _local_name(node.tag)
        if node_name in {"drawing", "pict", "object"}:
            image_count += 1
            continue
        if node_name == "extent":
            cx, cy = _extract_extent_inches(node)
            if cx > 0 and cy > 0:
                total_image_area += cx * cy

    return image_count, total_image_area


def _safe_document_body(document: Document) -> Any:
    try:
        return document.element.body
    except Exception:
        return None


def _collect_table_metrics(document: Document) -> Tuple[int, int]:
    table_row_count = 0
    word_count = 0

    for table in document.tables:
        for row in table.rows:
            table_row_count += 1
            for cell in row.cells:
                word_count += _count_words(cell.text)

    return table_row_count, word_count


def _estimate_pages_from_explicit_breaks(explicit_page_breaks: int) -> int:
    if explicit_page_breaks <= 0:
        return 0
    return explicit_page_breaks + 1


def _estimate_section_pages(document: Document) -> int:
    try:
        section_count = len(document.sections)
    except Exception:
        return 0

    if section_count > 1:
        return section_count
    return 0


def _estimate_pages_from_words(word_count: int) -> int:
    if word_count <= 0:
        return 0
    return max(1, math.ceil(word_count / DOCX_WORDS_PER_PAGE_ESTIMATE))


def _estimate_pages_from_blocks(block_count: int) -> int:
    if block_count <= 0:
        return 0
    return max(1, math.ceil(block_count / DOCX_BLOCKS_PER_PAGE_ESTIMATE))


def _estimate_docx_usable_page_area(document: Document) -> float:
    try:
        section = document.sections[0]
    except Exception:
        section = None

    if section is None:
        width = DOCX_DEFAULT_PAGE_WIDTH_INCH
        height = DOCX_DEFAULT_PAGE_HEIGHT_INCH
        left = right = top = bottom = DOCX_DEFAULT_MARGIN_INCH
    else:
        width = _safe_section_length_inches(
            section, "page_width", DOCX_DEFAULT_PAGE_WIDTH_INCH
        )
        height = _safe_section_length_inches(
            section, "page_height", DOCX_DEFAULT_PAGE_HEIGHT_INCH
        )
        left = _safe_section_length_inches(
            section, "left_margin", DOCX_DEFAULT_MARGIN_INCH
        )
        right = _safe_section_length_inches(
            section, "right_margin", DOCX_DEFAULT_MARGIN_INCH
        )
        top = _safe_section_length_inches(
            section, "top_margin", DOCX_DEFAULT_MARGIN_INCH
        )
        bottom = _safe_section_length_inches(
            section, "bottom_margin", DOCX_DEFAULT_MARGIN_INCH
        )

    usable_width = max(1.0, width - left - right)
    usable_height = max(1.0, height - top - bottom)
    return usable_width * usable_height


def _estimate_pages_from_images(
    image_count: int, total_image_area: float, usable_page_area: float
) -> int:
    if image_count <= 0:
        return 0

    area_based_estimate = 0
    if total_image_area > 0 and usable_page_area > 0:
        area_based_estimate = max(
            1,
            math.ceil(
                total_image_area / (usable_page_area * DOCX_IMAGE_PAGE_AREA_FACTOR)
            ),
        )

    image_count_based_estimate = 0
    if total_image_area > 0 and usable_page_area > 0:
        avg_image_area = total_image_area / image_count
        if avg_image_area >= usable_page_area * DOCX_IMAGE_DENSE_THRESHOLD:
            image_count_based_estimate = image_count
        elif avg_image_area >= DOCX_IMAGE_MEDIUM_AREA_INCH2:
            image_count_based_estimate = max(1, math.ceil(image_count / 2))
        elif avg_image_area >= DOCX_IMAGE_SMALL_AREA_INCH2:
            image_count_based_estimate = max(1, math.ceil(image_count / 3))
        else:
            image_count_based_estimate = max(1, math.ceil(image_count / 6))
    else:
        image_count_based_estimate = max(1, math.ceil(image_count / 6))

    return max(area_based_estimate, image_count_based_estimate)


def _extract_extent_inches(node: Any) -> Tuple[float, float]:
    # wp:extent stores width/height in EMU; values are used to estimate page occupancy.
    cx_emu = 0
    cy_emu = 0

    for key, value in getattr(node, "attrib", {}).items():
        key_name = str(key)
        if key_name.endswith("cx"):
            try:
                cx_emu = int(value)
            except Exception:
                cx_emu = 0
        elif key_name.endswith("cy"):
            try:
                cy_emu = int(value)
            except Exception:
                cy_emu = 0

    return _emu_to_inches(cx_emu, 0.0), _emu_to_inches(cy_emu, 0.0)


def _safe_section_length_inches(
    section: Any, attribute_name: str, default: float
) -> float:
    try:
        raw_value = getattr(section, attribute_name)
    except Exception:
        return default

    return _emu_to_inches(raw_value, default)


def _emu_to_inches(value: Any, default: float) -> float:
    try:
        numeric_value = float(value)
    except Exception:
        return default

    if numeric_value <= 0:
        return default

    return numeric_value / 914400.0


def _estimate_docx_pages_from_document_xml(document_xml: bytes) -> int:
    root = _parse_docx_document_root(document_xml)
    if root is None:
        return 0

    marker_counts, text_char_count = _collect_docx_xml_marker_counts(root)
    marker_estimate = _estimate_pages_from_docx_xml_markers(marker_counts)
    text_estimate = _estimate_pages_from_text_char_count(text_char_count)
    return max(marker_estimate, text_estimate)


def _parse_docx_document_root(document_xml: bytes) -> Optional[ET.Element]:
    try:
        root = ET.fromstring(document_xml)
    except Exception:
        return None

    if _local_name(root.tag) != "document":
        return None
    return root


def _collect_docx_xml_marker_counts(root: ET.Element) -> Tuple[dict, int]:
    marker_counts = {
        "manual_page_breaks": 0,
        "rendered_page_breaks": 0,
        "page_break_before_count": 0,
        "section_count": 0,
    }
    text_char_count = 0

    for node in root.iter():
        node_name = _local_name(node.tag)
        if node_name == "lastRenderedPageBreak":
            marker_counts["rendered_page_breaks"] += 1
        elif node_name == "br":
            if _is_page_break_node(node):
                marker_counts["manual_page_breaks"] += 1
        elif node_name == "pageBreakBefore":
            marker_counts["page_break_before_count"] += 1
        elif node_name == "sectPr":
            marker_counts["section_count"] += 1
        elif node_name == "t":
            text_char_count += _count_text_characters(node.text)

    return marker_counts, text_char_count


def _count_text_characters(text: Optional[str]) -> int:
    if not text:
        return 0
    return len(text.strip())


def _estimate_pages_from_docx_xml_markers(marker_counts: dict) -> int:
    marker_estimate = 0
    rendered_page_breaks = marker_counts["rendered_page_breaks"]
    manual_page_breaks = marker_counts["manual_page_breaks"]
    page_break_before_count = marker_counts["page_break_before_count"]
    section_count = marker_counts["section_count"]

    if rendered_page_breaks > 0:
        marker_estimate = max(marker_estimate, rendered_page_breaks + 1)
    if manual_page_breaks > 0:
        marker_estimate = max(marker_estimate, manual_page_breaks + 1)
    if page_break_before_count > 0:
        marker_estimate = max(marker_estimate, page_break_before_count + 1)
    if section_count > 1:
        marker_estimate = max(marker_estimate, section_count)
    return marker_estimate


def _estimate_pages_from_text_char_count(text_char_count: int) -> int:
    if text_char_count <= 0:
        return 0
    return max(1, math.ceil(text_char_count / DOCX_CHARS_PER_PAGE_ESTIMATE))


def _local_name(tag_name: str) -> str:
    if "}" in tag_name:
        return tag_name.rsplit("}", 1)[1]
    return tag_name


def check_doc_encrypted(uploaded_file: Any) -> Tuple[bool, Optional[str]]:
    if not is_ole_container(uploaded_file):
        return False, WORD_CORRUPT_ERROR

    try:
        uploaded_file.seek(0)
        head = uploaded_file.read(4096)
        uploaded_file.seek(0)
        if b"EncryptedPackage" in head or b"EncryptionInfo" in head:
            return False, WORD_PROTECTED_ERROR
    except Exception:
        return False, WORD_CORRUPT_ERROR

    return True, None


def check_doc_structure(uploaded_file: Any) -> Tuple[bool, Any]:
    if not is_ole_container(uploaded_file):
        return False, WORD_CORRUPT_ERROR

    try:
        uploaded_file.seek(0)
        content = uploaded_file.read(1024 * 1024)
        uploaded_file.seek(0)

        if b"WordDocument" not in content:
            return False, WORD_CORRUPT_ERROR

        return True, estimate_doc_page_count(content)
    except Exception:
        return False, WORD_CORRUPT_ERROR


def estimate_doc_page_count(content: bytes) -> int:
    """Best-effort page estimation for binary .doc content.

    Legacy .doc does not expose page count cheaply without full OLE parsing,
    so we estimate from page-break markers to make max-page checks enforceable.
    """
    if not content:
        return 0

    marker_count = content.count(b"\x0c")

    try:
        utf16_text = content.decode("utf-16-le", errors="ignore")
        marker_count = max(marker_count, utf16_text.count("\x0c"))
    except Exception:
        pass

    try:
        latin_text = content.decode("latin-1", errors="ignore")
        marker_count = max(marker_count, latin_text.count("\x0c"))
    except Exception:
        pass

    if marker_count <= 0:
        return 0

    return marker_count + 1


def check_word_page_count(page_count: int) -> Tuple[bool, Optional[str]]:
    if page_count > MAX_WORD_PAGES:
        return (
            False,
            f"Word exceeds the maximum allowed page count of {MAX_WORD_PAGES}.",
        )
    return True, None


def is_ole_container(uploaded_file: Any) -> bool:
    try:
        uploaded_file.seek(0)
        header = uploaded_file.read(len(OLE_SIGNATURE))
        uploaded_file.seek(0)
        return header == OLE_SIGNATURE
    except Exception:
        return False
